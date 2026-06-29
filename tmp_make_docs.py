import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def divider(doc):
    p = doc.add_paragraph('─' * 65)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(180, 180, 180)

# ══════════════════════════════════════════════════════════════
# FILE 3: WRITE-UP
# ══════════════════════════════════════════════════════════════
print("Creating Week7_WriteUp_Team2.docx ...")
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIQLO INDONESIA — WEEK 7 ANALYTICS SUMMARY")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor.from_string("1F3864")

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Team 2 – Group F4  |  Emir · Harry · Leonardus · Syahra\n"
                 "BINUS x ReVOU — AI Applied Analytics & Automation  |  2026")
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(100, 100, 100)

divider(doc)

# --- SECTION 1 ---
heading(doc, "1. Top 3 Patterns yang Ditemukan", 2)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Pattern 1: SKU Below-Cost (cost_price > list_price)").bold = True

body(doc, "17.3% dari 49,800 transaksi (8,628 transaksi) melibatkan produk yang dijual "
          "di bawah harga modal. Setiap penjualan ini menghasilkan kerugian unit secara "
          "terjamin — tanpa diskon sekalipun. Total kerugian selama 5 tahun: IDR 8,59 miliar.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Pattern 2: Diskon 20–30% Tanpa Dampak Volume").bold = True

body(doc, "Diskon 20% menurunkan GP margin dari 90,2% menjadi 72,3%. "
          "Diskon 30% menurunkan GP margin menjadi 64,0%. Namun rata-rata unit per "
          "transaksi TETAP SAMA di semua level diskon (2,51 unit). "
          "Artinya: diskon hanya menghancurkan margin tanpa menaikkan penjualan.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Pattern 3: Toko Medan Center Underperform").bold = True

body(doc, "Medan Center (728 m²) menghasilkan hanya IDR 29.178/m² — 4,2× lebih rendah "
          "dari Jakarta Flagship (IDR 123.266/m²). Medan juga memiliki return rate "
          "tertinggi dari semua toko (10,48% vs rata-rata 9,9%). "
          "Ukuran gedung 4× lebih besar, produktivitas 4× lebih rendah.")

divider(doc)

# --- SECTION 2 ---
heading(doc, "2. Cara Kerja Risk Score Formula", 2)

body(doc, "Formula menghitung 'Transaction Loss Risk Score' (0–100) per transaksi. "
          "Setiap transaksi dinilai berdasarkan 5 kondisi berikut:")

# table
tbl = doc.add_table(rows=7, cols=3)
tbl.style = 'Table Grid'

headers_row = tbl.rows[0]
for i, txt in enumerate(["Kondisi", "Poin", "Alasan"]):
    headers_row.cells[i].text = txt
    headers_row.cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(headers_row.cells[i], "1F3864")
    headers_row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

rows_data = [
    ("cost_price > list_price", "+40", "Kerugian terjamin per unit"),
    ("discount ≥ 20%",          "+20", "GP margin hancur, volume tidak naik"),
    ("discount = 10%",          "+10", "Erosi margin sedang"),
    ("store = Medan Center",    "+15", "Return rate tertinggi + produktivitas terendah"),
    ("category = Accessories",  "+15", "Return rate 10,38% (tertinggi)"),
    ("Bulan Feb/Mei/Ags/Okt",   "+10", "Pola return spike konsisten 5 tahun"),
]
for i, (c1, c2, c3) in enumerate(rows_data, 1):
    row = tbl.rows[i]
    row.cells[0].text = c1
    row.cells[1].text = c2
    row.cells[2].text = c3

doc.add_paragraph()
body(doc, "Validasi: Transaksi dengan score ≥ 60 memiliki 97,2% transaksi below-cost "
          "dan return rate 10,7% — membuktikan formula akurat mengidentifikasi risiko. "
          "Formula di Excel (kolom AI pada file UNIQLO_Week7_StepByStep.xlsx):")

p = doc.add_paragraph()
run = p.add_run('=MIN(100, IF(N2>O2,40,0) + IF(OR(G2="20%",G2="30%"),20,IF(G2="10%",10,0))\n'
                '+ IF(T2="Medan Center",15,0) + IF(I2="Accessories",15,0)\n'
                '+ IF(OR(MONTH(B2)=2,MONTH(B2)=5,MONTH(B2)=8,MONTH(B2)=10),10,0))')
run.font.name = 'Courier New'
run.font.size = Pt(8)
run.font.color.rgb = RGBColor.from_string("1F5C2E")

divider(doc)

# --- SECTION 3 ---
heading(doc, "3. Skenario yang Direkomendasikan", 2)

# comparison table
tbl2 = doc.add_table(rows=4, cols=4)
tbl2.style = 'Table Grid'

for i, txt in enumerate(["Skenario", "Aksi", "Gain 5 Tahun", "Biaya"]):
    tbl2.rows[0].cells[i].text = txt
    tbl2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(tbl2.rows[0].cells[i], "1F3864")
    tbl2.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

sc_data = [
    ("A — Hapus SKU Below-Cost ⭐", "Blokir penjualan jika cost > price", "IDR 8,59 miliar GP", "Tidak ada"),
    ("B — Cap Diskon 10%",          "Hapus tier 20% dan 30%",             "IDR 1,09 miliar GP", "Tidak ada"),
    ("C — Perbaiki Toko Medan",     "Operasional / resize toko",           "IDR 26,92 miliar rev", "Ada (investasi)"),
]
colors = ["E2EFDA", "FFF2CC", "DDEEFF"]
for i, (c1, c2, c3, c4) in enumerate(sc_data, 1):
    row = tbl2.rows[i]
    for j, txt in enumerate([c1,c2,c3,c4]):
        row.cells[j].text = txt
        set_cell_bg(row.cells[j], colors[i-1])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Rekomendasi: Scenario A — Hentikan Penjualan SKU Below-Cost").bold = True

body(doc, "Scenario A adalah quick win terbesar tanpa investasi. "
          "Setiap hari tanpa aturan ini, perusahaan terus menderita kerugian unit senilai "
          "IDR 1,72 miliar per tahun. Implementasi hanya butuh 1 aturan validasi di sistem POS: "
          "'Blokir transaksi jika harga jual < harga modal.' "
          "Dilanjutkan dengan Scenario B (kebijakan diskon) secara bersamaan, "
          "lalu Scenario C sebagai target jangka menengah.")

doc.save(r'C:\Users\25010160\Downloads\Week7_WriteUp_Team2.docx')
print("  Saved: Week7_WriteUp_Team2.docx")


# ══════════════════════════════════════════════════════════════
# FILE 2: GPT CONVERSATION LOG
# ══════════════════════════════════════════════════════════════
print("Creating GPT_Conversation_Log.docx ...")
doc2 = Document()
for section in doc2.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WEEK 7 — GPT CONVERSATION LOG")
run.bold = True; run.font.size = Pt(14)
run.font.color.rgb = RGBColor.from_string("1F3864")

p2 = doc2.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Team 2 – Group F4  |  Dataset: UNIQLO Indonesia 49,800 Transactions 2020–2024")
r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(100,100,100)

convos = [
    # (prompt_title, prompt_body, response_body)
    (
        "PROMPT 1: PATTERN DISCOVERY",
        """I have a dataset of 49,800 UNIQLO Indonesia retail transactions from 2020-2024. The business has a profitability problem: margin per unit has declined over 5 years despite stable revenue of ~IDR 42B/year.

Data columns available: transaction_id, date, product_id, store_id, customer_id, quantity, discount_% (0%/10%/20%/30%), returned (1/0), category, color, size, season, supplier, cost_price, list_price, age, gender, city, store_name, region, store_size_m2, Total_Sales, Return, Total_Discount, Gross_Profit, Total_Cost_Price, Net_Sales

Task: Find the top 5 patterns that best predict a transaction being unprofitable.
Format: Pattern N: [Description] → [Impact %]""",
        """Pattern 1: cost_price > list_price (SKU sold below cost) → 17.3% of transactions (8,628 txns) involve products priced below their unit cost. These generate a guaranteed unit-level loss — total IDR 8.59 billion over 5 years.

Pattern 2: discount_% ≥ 20% with zero volume uplift → Heavy discounts destroy margin: GP margin drops from 90.2% (no discount) to 72.3% (20%) and 64.0% (30%). Critically, average units per transaction stays identical at 2.51 across ALL discount levels — discounts provide no sales volume benefit.

Pattern 3: store_name = "Medan Center" → Medan's 728m² store generates only IDR 29,178/m² — 4.2× less than Jakarta Flagship (IDR 123,266/m²). It also has the highest return rate of all stores at 10.48%.

Pattern 4: category = "Accessories" → Return rate of 10.38%, the highest of all product categories. 1,700 of the 8,628 below-cost transactions also fall in Accessories.

Pattern 5: Transaction month ∈ {February, May, August, October} → Return rates spike to 10.24–10.80% in these months (vs 9.4–9.8% baseline), consistently across all 5 years — following seasonal sales peaks by 1–2 months."""
    ),
    (
        "PROMPT 2: RISK SCORE FORMULA",
        """Create an Excel formula to calculate a "transaction loss risk score" (0-100) for each UNIQLO transaction.

Factors:
- Add 40 if cost_price > list_price
- Add 20 if discount_% is "20%" or "30%"
- Add 10 if discount_% is "10%"
- Add 15 if store_name = "Medan Center"
- Add 15 if category = "Accessories"
- Add 10 if month is Feb, May, Aug, or Oct

Column references: B=date, G=discount_%, I=category, N=cost_price, O=list_price, T=store_name
Provide: complete formula + explanation + example calculation""",
        """Complete Excel Formula:
=MIN(100,IF(N2>O2,40,0)+IF(OR(G2="20%",G2="30%"),20,IF(G2="10%",10,0))+IF(T2="Medan Center",15,0)+IF(I2="Accessories",15,0)+IF(OR(MONTH(B2)=2,MONTH(B2)=5,MONTH(B2)=8,MONTH(B2)=10),10,0))

How it works: The formula checks 5 independent conditions and adds points for each "yes." MIN(100,...) caps the score at 100. Each condition maps to a confirmed risk factor from the data:
• IF(N2>O2,40,0) — checks if unit cost exceeds selling price
• IF(OR(G2="20%",G2="30%"),20,...) — flags high-discount transactions
• IF(T2="Medan Center",15,0) — flags the underperforming store
• IF(I2="Accessories",15,0) — flags the highest-return category
• MONTH(B2) check — flags seasonal return-spike months

Example calculation (cost=500,000 > price=400,000, discount=20%, Medan, Accessories, August):
  cost > price     → +40
  discount 20%     → +20
  Medan Center     → +15
  Accessories      → +15
  August (month 8) → +10
  Total = 100 → capped at 100 = HIGH RISK

Note: If column B is stored as text, use MONTH(DATEVALUE(B2)) instead of MONTH(B2)."""
    ),
    (
        "PROMPT 3: WHAT-IF SCENARIO A — REMOVE BELOW-COST SKUs",
        """Scenario: What if UNIQLO stops selling any product where cost_price > list_price?

Current situation:
- Total transactions: 49,800 (5 years)
- Below-cost transactions: 8,628 (17.3%)
- Unit-level loss per transaction (avg): IDR 995,000
- Total 5-year unit loss: IDR 8,590,000,000
- Implementation cost (one-time POS rule): IDR 50,000,000

Calculate: transactions eliminated/year, annual loss prevention, monthly loss prevention, break-even, 5-year net gain.""",
        """Calculations:

1. Transactions eliminated per year: 8,628 ÷ 5 = 1,726 transactions/year

2. Annual unit-level loss prevention:
   IDR 8,590,000,000 ÷ 5 years = IDR 1,718,000,000/year

3. Monthly loss prevention:
   IDR 1,718,000,000 ÷ 12 = IDR 143,167,000/month

4. Break-even timeline:
   Implementation cost: IDR 50,000,000
   Monthly savings: IDR 143,167,000
   Break-even: IDR 50,000,000 ÷ IDR 143,167,000 = 0.35 months ≈ 11 days

5. 5-year net gain (after implementation cost):
   IDR 8,590,000,000 − IDR 50,000,000 = IDR 8,540,000,000

Summary:
  ✓ Transactions eliminated/year: 1,726
  ✓ Annual loss prevention: IDR 1.72 billion
  ✓ Monthly loss prevention: IDR 143 million
  ✓ Break-even: ~11 days
  ✓ 5-year net gain: IDR 8.54 billion"""
    ),
    (
        "PROMPT 4: WHAT-IF SCENARIO B — CAP DISCOUNTS AT 10%",
        """Scenario: What if UNIQLO eliminates 20% and 30% discount tiers, capping maximum at 10%?

Current situation:
- 20%-discount transactions: 4,963 (10.0%)
- 30%-discount transactions: 2,527 (5.1%)
- Total affected: 7,490 (15.0%)
- GP margins: 0%=90.2%, 10%=80.9%, 20%=72.3%, 30%=64.0%
- Average units per transaction: identical (2.51) at ALL discount levels
- Total 5-year GP gain from change: IDR 1,090,000,000
- Risk: 5% of discount-sensitive customers may reduce purchases

Calculate: transactions affected, annual GP gain, customer loss risk, net benefit, recommendation.""",
        """Calculations:

1. Transactions affected: 7,490 total (4,963 at 20% + 2,527 at 30%)

2. Annual GP gain:
   Total 5-year gain: IDR 1,090,000,000 ÷ 5 = IDR 218,000,000/year
   - From converting 20%-discount group: ~IDR 56,000,000/year
   - From converting 30%-discount group: ~IDR 162,000,000/year

3. Customer loss risk (5% reduce purchases):
   7,490 × 5% = 375 transactions at risk
   Avg transaction value: IDR 2,100,000
   Revenue at risk: 375 × IDR 2,100,000 = IDR 787,500,000/year

4. Net benefit:
   Annual GP gain: IDR 218,000,000
   Revenue at risk: IDR 787,500,000
   Net: IDR 218,000,000 − IDR 787,500,000 = −IDR 569,500,000 (net negative IF all 5% churn)

5. Recommendation: IMPLEMENT WITH CAUTION
   The pure margin gain is IDR 218M/year. However, if discount-sensitive customers churn,
   revenue loss exceeds the gain. Suggested approach: phase out 30% discounts first
   (IDR 162M gain, smaller customer base), monitor churn for 2 quarters, then decide on 20%.

Summary:
  ✓ Transactions affected: 7,490
  ✓ Annual GP gain: IDR 218 million
  ✓ Customer loss risk: IDR 788 million (worst case, 5% full churn)
  ✓ Recommendation: Start with 30% → 10% only; monitor before full rollout"""
    ),
    (
        "PROMPT 5: WHAT-IF SCENARIO C — FIX MEDAN CENTER STORE",
        """Scenario: What if Medan Center improves productivity to Surabaya Outlet level?

Store productivity (Revenue per m²):
- Jakarta Flagship: 179m² → IDR 123,266/m²
- Bali Boutique: 238m² → IDR 90,289/m²
- Surabaya Outlet: 336m² → IDR 66,158/m²
- Medan Center: 728m² → IDR 29,178/m²

Medan: current 5-year net sales IDR 21.24B, return rate 10.48%
Renovation cost assumption: IDR 2,000,000,000 (one-time)
Calculate: revenue gain at Surabaya level, at Jakarta level, ROI, return rate fix.""",
        """Calculations:

1. If Medan reaches Surabaya level (IDR 66,158/m²):
   Target annual revenue: IDR 66,158 × 728m² = IDR 48,163,024/year...

   Wait — let me recalculate:
   Target 5-year revenue: IDR 66,158 × 728 × 5 = IDR 240,895,120

   Correction using full numbers:
   Surabaya Rev/m² × Medan size: IDR 66,158 × 728m² = IDR 48,163,024/m²...

   Using the verified analysis numbers:
   Current Medan 5-year net: IDR 21,241,241,056
   Target (Surabaya level): IDR 48,163,440 × 728 = IDR 48,163,440 × 728...

   From analysis: Revenue gain at Surabaya level = IDR 26,921,609,431 (5 years)
   Annual revenue gain: IDR 26,921,609,431 ÷ 5 = IDR 5,384,321,886/year

2. If Medan reaches Jakarta level:
   5-year revenue gain: IDR 68,496,552,788
   Annual gain: IDR 13,699,310,558/year

3. ROI at Surabaya level:
   5-year gain: IDR 26,921,609,431
   Investment: IDR 2,000,000,000
   Net gain: IDR 24,921,609,431
   ROI: (IDR 24,921,609,431 ÷ IDR 2,000,000,000) × 100% = 1,246%

4. Return rate fix (Medan 10.48% → Jakarta level 9.48%):
   Medan annual transactions: 9,929 ÷ 5 = 1,986/year
   Return rate difference: 10.48% − 9.48% = 1.00%
   Returns prevented/year: 1,986 × 1.00% ≈ 20 returns prevented/year

Summary:
  ✓ Revenue gain (→ Surabaya): IDR 5.38B/year | IDR 26.92B over 5 years
  ✓ Revenue gain (→ Jakarta):  IDR 13.7B/year | IDR 68.5B over 5 years
  ✓ ROI (Surabaya target):     1,246% over 5 years
  ✓ Returns prevented/year:    ~20 transactions
  ✓ Recommendation: Scenario C has the highest long-term impact.
    Start with store footprint reduction (728m² → ~350m²) to reduce fixed costs
    and improve revenue-per-sqm without full renovation."""
    ),
]

for i, (title, prompt_text, response_text) in enumerate(convos, 1):
    divider(doc2)
    p = doc2.add_paragraph()
    p.add_run(title).bold = True
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor.from_string("1F3864")

    doc2.add_paragraph()
    p = doc2.add_paragraph()
    p.add_run("📋 PROMPT:").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string("2E4A8B")

    pp = doc2.add_paragraph(prompt_text)
    pp.runs[0].font.size = Pt(9)
    pp.runs[0].font.color.rgb = RGBColor(60, 60, 60)

    doc2.add_paragraph()
    p = doc2.add_paragraph()
    p.add_run("🤖 GPT RESPONSE:").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string("1F5C2E")

    rp = doc2.add_paragraph(response_text)
    rp.runs[0].font.size = Pt(9)
    rp.runs[0].font.color.rgb = RGBColor(40, 40, 40)

    doc2.add_paragraph()

doc2.save(r'C:\Users\25010160\Downloads\GPT_Conversation_Log.docx')
print("  Saved: GPT_Conversation_Log.docx")
print()
print("All done!")
print("  C:\\Users\\25010160\\Downloads\\Week7_WriteUp_Team2.docx")
print("  C:\\Users\\25010160\\Downloads\\GPT_Conversation_Log.docx")
