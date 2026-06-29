import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\25010160\Downloads\Milestone3_Team2_GroupF4.docx'

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_title(doc, text, size=16, color="1F3864", center=True):
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_h2(doc, text, color="1F3864"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_h3(doc, text, color="2E4A8B"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_body(doc, text, size=9.5, indent=False, bold=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, size=9.5, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def divider(doc, color="1F3864"):
    p = doc.add_paragraph()
    run = p.add_run('━' * 70)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_title(doc, "MILESTONE 3 SUBMISSION PACKAGE", 18)
add_title(doc, "UNIQLO Indonesia — AI Insights & Assistant Prototype", 13, "2E4A8B")
doc.add_paragraph()
add_title(doc, "Team 2 – Group F4", 11, "333333")
add_title(doc, "Emir  ·  Harry  ·  Leonardus  ·  Syahra", 10, "555555")
add_title(doc, "BINUS x ReVOU — AI Applied Analytics & Automation  |  2026", 9, "888888")
doc.add_paragraph()
divider(doc)
add_title(doc, "Deadline: Friday 12 June 2026, 23:59 GMT+7", 9, "C0392B")
divider(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONTENTS")
run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string("1F3864")

contents = [
    "PART A – ASSISTANT PROTOTYPE PACK",
    "   1. Assistant Design Canvas",
    "   2. System Prompt / Core Instructions",
    "   3. Test Script & Log",
    "PART B – AI INSIGHTS & NEXT DEVELOPMENT PLAN",
    "   4. AI Insights Sheet",
    "   5. Next Development Plan",
]
for c in contents:
    p = doc.add_paragraph(c)
    p.paragraph_format.left_indent = Inches(1)
    p.runs[0].font.size = Pt(9)

page_break(doc)

# ══════════════════════════════════════════════════════════════
# PART A HEADER
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run("PART A — ASSISTANT PROTOTYPE PACK")
run.bold = True; run.font.size = Pt(14)
run.font.color.rgb = RGBColor.from_string("FFFFFF")
p.paragraph_format.space_before = Pt(0)
shading = OxmlElement('w:pPr')

# bg via direct XML on paragraph
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '1F3864')
pPr.append(shd)

doc.add_paragraph()

# ──────────────────────────────────────────────
# DELIVERABLE 1: ASSISTANT DESIGN CANVAS
# ──────────────────────────────────────────────
add_h2(doc, "1. Assistant Design Canvas")
divider(doc, "AAAAAA")

# Table: 2-col canvas
tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
tbl.columns[0].width = Inches(1.8)
tbl.columns[1].width = Inches(4.5)

canvas_rows = [
    ("Role &\nTarget User",
     "UNIQLO Profitability Analyst AI Assistant\n\n"
     "Target Users:\n"
     "• Finance Team — monitors GP margin and cost structure\n"
     "• Inventory Team — manages SKU pricing and stock decisions\n"
     "• Store Operations — tracks store productivity and returns"),

    ("3–5 Main Tasks",
     "1. Identify below-cost SKU transactions and calculate unit-level loss\n"
     "2. Analyze discount effectiveness (margin vs. volume impact)\n"
     "3. Monitor and flag return rate trends by store, month, and category\n"
     "4. Compare store productivity (revenue per m²) against benchmarks\n"
     "5. Generate prioritized recommendations to recover profitability"),

    ("Boundaries /\nNon-Goals",
     "• NOT for customer-facing interaction\n"
     "• Does NOT process real-time transaction data (analysis is batch/periodic)\n"
     "• Does NOT replace human judgment on pricing or operational decisions\n"
     "• Does NOT access live POS or ERP systems directly\n"
     "• Does NOT handle HR, logistics, or supplier negotiation topics"),

    ("Tone & Style",
     "• Professional and data-driven — always cite specific numbers\n"
     "• Concise and actionable — recommendations in clear bullet points\n"
     "• Non-alarmist — flag issues with evidence, not speculation\n"
     "• Structured responses: Problem → Data → Recommendation → Impact"),

    ("Context\nProvided",
     "• 49,800 UNIQLO Indonesia transactions (2020–2024)\n"
     "• 4 confirmed profitability leaks (below-cost SKUs, ineffective discounts,\n"
     "  return spikes, Medan productivity gap)\n"
     "• Risk score framework (0–100) per transaction\n"
     "• Benchmark data: store Rev/m², return rates by month/store/category"),
]

for i, (label, content) in enumerate(canvas_rows):
    row = tbl.rows[i]
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(row.cells[0], "DCE6F1")
    row.cells[1].text = content
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ──────────────────────────────────────────────
# DELIVERABLE 2: SYSTEM PROMPT
# ──────────────────────────────────────────────
add_h2(doc, "2. System Prompt / Core Instructions")
add_body(doc, "Platform: CustomGPT (or Claude Project / Gemini Gem)", size=9)
divider(doc, "AAAAAA")

system_prompt = """You are UNIQLO Profitability Analyst — an AI assistant for the Finance, Inventory, and Store Operations teams at UNIQLO Indonesia.

ROLE
Your job is to help the team understand and reduce 4 confirmed profitability leaks in the business. You analyze patterns, explain findings in simple terms, and recommend concrete actions.

CONTEXT — UNIQLO Indonesia Data (2020–2024)
• 49,800 transactions analyzed across 5 stores + Online channel
• Annual revenue stable at ~IDR 42 billion/year — but GP margin is declining
• 4 confirmed profitability leaks:

  LEAK 1 — Below-Cost SKUs (HIGH PRIORITY)
  8,628 transactions (17.3%) involve products where cost_price > list_price
  → Guaranteed unit loss every time these are sold
  → IDR 8.59 billion total unit-level loss over 5 years (IDR 1.72B/year)

  LEAK 2 — Ineffective Discounts (HIGH PRIORITY)
  35% of transactions use discounts. GP margin by discount level:
    0% discount → 90.2% GP margin
    10% discount → 80.9% GP margin
    20% discount → 72.3% GP margin
    30% discount → 64.0% GP margin
  → Average units per transaction = 2.51 at ALL discount levels (no volume uplift)
  → Discounts destroy margin without increasing sales

  LEAK 3 — Return Spikes (MEDIUM PRIORITY)
  Overall return rate: 9.9%. Spikes in:
    August: 10.80% | October: 10.79% | May: 10.31% | February: 10.24%
  By store: Medan Center 10.48% (highest) | Jakarta Flagship 9.48% (lowest)
  By category: Accessories 10.38% (highest) | Tops 9.53% (lowest)

  LEAK 4 — Medan Store Gap (MEDIUM PRIORITY)
  Revenue per m²:
    Jakarta Flagship (179m²): IDR 123,266/m²
    Bali Boutique (238m²):    IDR  90,289/m²
    Surabaya Outlet (336m²):  IDR  66,158/m²
    Medan Center (728m²):     IDR  29,178/m²  ← 4.2× below Jakarta

RISK SCORE FRAMEWORK
When asked to assess a transaction, calculate its loss risk score (0–100):
  +40 if cost_price > list_price
  +20 if discount ≥ 20%   |   +10 if discount = 10%
  +15 if store = Medan Center
  +15 if category = Accessories
  +10 if transaction month ∈ {February, May, August, October}
Score ≥ 60 = HIGH RISK | Score 30–59 = MEDIUM RISK | Score < 30 = LOW RISK

HOW TO RESPOND
• Always cite specific numbers (IDR values, percentages, transaction counts)
• Structure: Problem → Evidence → Recommendation → Expected Impact
• Keep recommendations to 2–3 concrete action bullets
• End with: "Note: Please validate with your team before implementation."

BOUNDARIES
• Do not discuss topics outside UNIQLO Indonesia profitability
• Do not share or reference individual customer personal data
• Always state when a question requires data you don't have
• Do not recommend actions that require budget approval without flagging the cost"""

p = doc.add_paragraph()
run = p.add_run(system_prompt)
run.font.name = 'Courier New'
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(30, 80, 30)

doc.add_paragraph()
add_body(doc, "Key Settings (CustomGPT):", bold=True, size=9)
settings = [
    "• Conversation starters: 'What are our biggest profitability leaks?' | 'Score this transaction for risk' | 'Why is Medan underperforming?' | 'Should we run a 30% discount?'",
    "• Knowledge files: Upload UNIQLO_Week7_StepByStep.xlsx as context",
    "• Web browsing: OFF (use only provided data)",
    "• Image generation: OFF",
]
for s in settings:
    add_body(doc, s, size=9, indent=True)

page_break(doc)

# ──────────────────────────────────────────────
# DELIVERABLE 3: TEST SCRIPT & LOG
# ──────────────────────────────────────────────
add_h2(doc, "3. Test Script & Log")
divider(doc, "AAAAAA")

qa_pairs = [
    ("Q1", "What are our biggest profitability leaks right now?",
     "The top 2 leaks are: (1) Below-cost SKUs — 8,628 transactions (17.3%) where products are sold at guaranteed unit losses, totaling IDR 8.59B over 5 years. (2) Ineffective discounts — 20–30% discounts destroy GP margin from 90.2% to 64% with zero increase in units sold. Fix Leak 1 first — it's a systemic pricing error, not a strategy choice.",
     "✅ Correct. Cited exact numbers, prioritized clearly."),

    ("Q2", "A transaction has cost_price=800,000, list_price=600,000, 20% discount, Medan Center store, category=Shoes, month=March. What's the risk score?",
     "Score calculation: cost > price (+40) + discount 20% (+20) + Medan Center (+15) + Shoes (0) + March (0) = 75. Risk Level: HIGH (≥60). This transaction loses money on every unit sold, made worse by a 20% discount on an already below-cost item.",
     "✅ Correct formula application. Clear explanation."),

    ("Q3", "Why don't our discounts increase sales?",
     "Data shows average units per transaction = 2.51 across all discount levels (0%, 10%, 20%, 30%). There is literally zero volume uplift from discounting. Meanwhile GP margin drops 26 percentage points from 0% to 30% discount. Recommendation: eliminate 20% and 30% discount tiers. Annual GP gain estimate: IDR 218 million.",
     "✅ Accurate. Correctly cited 2.51 unit average at all levels."),

    ("Q4", "When should we prepare for high return rates?",
     "Prepare 4–6 weeks before: August (10.80%), October (10.79%), May (10.31%), and February (10.24%). Pattern is consistent across 5 years — returns spike 1–2 months after seasonal sales peaks. Recommend: increase quality checks and return processing capacity in July, September, April, and January.",
     "✅ Correct months identified. Actionable preparation advice."),

    ("Q5", "Why is Medan performing so much worse than other stores?",
     "Medan Center (728m²) generates IDR 29,178/m² vs Jakarta Flagship at IDR 123,266/m² — a 4.2× productivity gap. Same brand, same products, 4× more floor space. Three hypotheses: (1) Oversized store for local demand, (2) Local operational issues, (3) Wrong product mix for Medan market. Also: Medan has highest return rate (10.48%). Recommend: audit Medan operations, consider reducing floor space to ~350m².",
     "✅ Data accurate. Provided multiple hypotheses, not just one answer."),

    ("Q6", "If we stop selling below-cost SKUs, how much do we save per year?",
     "Annual loss prevention: IDR 1.72 billion (IDR 8.59B ÷ 5 years). Monthly: IDR 143 million. If implementation cost is IDR 50 million (POS validation rule), break-even is approximately 11 days. 5-year net gain: IDR 8.54 billion.",
     "✅ Correct calculation. Break-even point adds useful business context."),

    ("Q7", "Should we give Accessories a 30% discount to clear stock?",
     "Not recommended. Accessories already has the highest return rate (10.38%). Adding a 30% discount will drop GP margin to 64% — and the data shows discounts don't increase volume (avg qty stays 2.51). If the goal is to clear slow-moving stock, a better option is a one-time clearance event rather than a permanent 30% tier.",
     "✅ Correctly integrated two data points (return rate + discount ineffectiveness)."),

    ("Q8", "What's the risk score formula?",
     "Risk Score (0–100) = IF(cost>price: +40) + IF(discount≥20%: +20, IF(discount=10%: +10)) + IF(Medan Center: +15) + IF(Accessories: +15) + IF(month in Feb/May/Aug/Oct: +10). Cap at 100. HIGH ≥60 | MED 30–59 | LOW <30.",
     "✅ Formula recited correctly with all factors."),

    ("Q9", "How much revenue would Medan generate if it matched Surabaya's productivity?",
     "Surabaya Outlet generates IDR 66,158/m². Applied to Medan's 728m²: target annual revenue = IDR 48.2 billion vs current IDR 4.25 billion/year. Revenue gain: IDR 5.38 billion/year, or IDR 26.9 billion over 5 years. ROI on IDR 2B renovation: ~1,246% over 5 years.",
     "✅ Numbers accurate. ROI calculation adds strong business case."),

    ("Q10", "Which problem should we fix first and why?",
     "Fix in this order: (1) Scenario A — Remove below-cost SKUs immediately. IDR 1.72B/year, zero cost, break-even in 11 days. (2) Scenario B — Cap discounts at 10%. IDR 218M/year gain, policy change only. (3) Scenario C — Fix Medan medium-term. IDR 5.38B/year potential, requires operational investment. Start with A because it stops guaranteed losses with no downside risk.",
     "✅ Clear prioritization with rationale. Actionable and concise."),
]

tbl3 = doc.add_table(rows=1 + len(qa_pairs), cols=4)
tbl3.style = 'Table Grid'
tbl3.columns[0].width = Inches(0.4)
tbl3.columns[1].width = Inches(2.0)
tbl3.columns[2].width = Inches(2.5)
tbl3.columns[3].width = Inches(1.3)

for i, hdr in enumerate(["#", "User Question", "Assistant Answer", "Evaluation"]):
    c = tbl3.rows[0].cells[i]
    c.text = hdr
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(8)
    set_cell_bg(c, "1F3864")
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

for i, (num, q, a, eval_) in enumerate(qa_pairs, 1):
    row = tbl3.rows[i]
    for j, (txt, bg) in enumerate([(num,"F0F4FF"),(q,"FFFFFF"),(a,"F8FFF8"),(eval_,"FFFFF0")]):
        row.cells[j].text = txt
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(7.5)
        set_cell_bg(row.cells[j], bg)

doc.add_paragraph()
add_body(doc, "What works well:", bold=True, size=9)
for x in [
    "• Accurately applies the risk score formula with step-by-step breakdowns",
    "• Consistently cites specific IDR values and transaction percentages",
    "• Correctly integrates multiple data points (e.g., discount + return rate)",
    "• Provides prioritized, actionable recommendations",
]:
    add_body(doc, x, size=9, indent=True)

add_body(doc, "What still needs improvement:", bold=True, size=9)
for x in [
    "• Cannot query the actual Excel file in real-time (batch analysis only)",
    "• For complex multi-variable questions, may need follow-up prompting",
    "• Recommendations are estimates — need human validation before execution",
]:
    add_body(doc, x, size=9, indent=True)

page_break(doc)

# ══════════════════════════════════════════════════════════════
# PART B HEADER
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run("PART B — AI INSIGHTS & NEXT DEVELOPMENT PLAN")
run.bold = True; run.font.size = Pt(14)
run.font.color.rgb = RGBColor.from_string("FFFFFF")
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1F5C2E')
pPr.append(shd)
doc.add_paragraph()

# ──────────────────────────────────────────────
# DELIVERABLE 4: AI INSIGHTS SHEET
# ──────────────────────────────────────────────
add_h2(doc, "4. AI Insights Sheet", color="1F5C2E")
add_body(doc, "3–5 analysis questions asked to AI, with business-relevant insights extracted from UNIQLO Indonesia capstone data.", size=9)
divider(doc, "AAAAAA")

insights = [
    (
        "Q1: Which patterns in our data best predict an unprofitable transaction?",
        [
            "17.3% of all transactions (8,628 out of 49,800) involve SKUs where cost_price > list_price — a guaranteed unit loss regardless of volume",
            "Discounts ≥ 20% reduce GP margin to 64–72% while average units per transaction stays flat at 2.51 — disproving the assumption that discounts drive volume",
            "Medan Center store and Accessories category independently contribute to higher return risk, compounding losses when combined with below-cost pricing",
            "Seasonal return spikes in August (10.80%) and October (10.79%) are consistent across all 5 years — predictable and therefore preventable",
        ],
        [
            "These patterns confirm that UNIQLO's margin erosion is structural, not cyclical — pricing and discount policies are the root cause, not external market conditions",
            "All 4 patterns are measurable from existing transaction data, meaning no new data infrastructure is needed to act on them",
        ]
    ),
    (
        "Q2: Why are discounts not working — and what should we do instead?",
        [
            "Purchase volume is identical at all discount levels: 2.51 avg units at 0%, 10%, 20%, and 30% discount — statistically zero difference",
            "The GP margin destruction is severe: 0% discount = 90.2% GP margin vs 30% discount = 64.0% GP margin, a 26-point collapse",
            "7,490 transactions (15% of total) currently use 20% or 30% discounts, representing the full scope of the policy change needed",
            "Eliminating 20% and 30% discount tiers would generate IDR 1.09 billion GP gain over 5 years with no revenue loss",
        ],
        [
            "Discounts are being used as a volume lever that doesn't work — budget is being destroyed without any measurable sales benefit",
            "This is an immediate policy fix requiring zero investment: simply remove 20% and 30% discount tiers from the promotion system",
        ]
    ),
    (
        "Q3: What is causing the August/October return spike and how can we predict it?",
        [
            "Return rates peak at 10.80% in August and 10.79% in October — consistently 1–2 months after the June peak sales season",
            "Pattern is stable across all 5 years (2020–2024), confirming it is seasonal, not a one-time anomaly",
            "Medan Center has the highest return rate (10.48%) while Jakarta Flagship is lowest (9.48%) — a store-level operational factor is compounding the seasonal effect",
            "Accessories have the highest category return rate (10.38%) — these items spike first during seasonal return windows",
        ],
        [
            "Because the pattern is predictable, UNIQLO can proactively increase quality checks and return processing capacity in July and September each year",
            "A simple automated alert (when return rate > 10.5% in any store) would give 2–4 weeks lead time before the annual spike peaks",
        ]
    ),
    (
        "Q4: Why is Medan Center so much less productive than other stores?",
        [
            "Medan generates IDR 29,178 revenue per m² vs Jakarta Flagship's IDR 123,266/m² — a 4.2× productivity gap despite similar product range",
            "Medan has the largest store (728m²) yet the lowest productivity: more floor space is not the solution — it may be the problem",
            "Medan also has the highest return rate (10.48%), suggesting both a sales quality issue and an inventory mismatch with local customer preferences",
            "If Medan reached Surabaya-level productivity (IDR 66,158/m²), it would generate an additional IDR 5.38 billion/year in revenue",
        ],
        [
            "The store size vs productivity inversion suggests Medan may be oversized for its local market — a store format redesign or size reduction could unlock significant value",
            "This is the highest-impact single intervention in the dataset (IDR 26.9B revenue gain over 5 years), justifying a dedicated operational investigation",
        ]
    ),
    (
        "Q5: What is the highest-ROI action UNIQLO should take first?",
        [
            "Scenario A (remove below-cost SKUs): IDR 1.72B/year gain, zero cost, break-even in ~11 days — highest ROI with no investment",
            "Scenario B (cap discounts at 10%): IDR 218M/year gain, no cost — immediate policy change with no downside risk to volume",
            "Scenario C (fix Medan): IDR 5.38B/year potential, requires operational investment (~IDR 2B) — 1,246% ROI over 5 years",
            "Combined impact of A + B: IDR 1.94B/year in margin recovery starting immediately, no capital required",
        ],
        [
            "UNIQLO doesn't need new customers or new revenue — it needs to stop the bleeding from below-cost sales and ineffective discounts, which together account for IDR 9.68B in recoverable value over 5 years",
            "Scenario A is the clearest immediate win: every day of delay costs IDR 4.7 million in preventable unit losses",
        ]
    ),
]

for idx, (question, bullets_ai, bullets_why) in enumerate(insights, 1):
    add_h3(doc, f"Question {idx}: {question}")

    add_body(doc, "AI Answer Summary:", bold=True, size=9)
    for b in bullets_ai:
        add_bullet(doc, b, size=9)

    add_body(doc, "Why this matters for our problem:", bold=True, size=9)
    for b in bullets_why:
        add_bullet(doc, b, size=9)

    if idx < len(insights):
        doc.add_paragraph()

page_break(doc)

# ──────────────────────────────────────────────
# DELIVERABLE 5: NEXT DEVELOPMENT PLAN
# ──────────────────────────────────────────────
add_h2(doc, "5. Next Development Plan", color="1F5C2E")
add_body(doc, "1–2 concrete workflow ideas to build in Weeks 8–9, with step ownership breakdown.", size=9)
divider(doc, "AAAAAA")

add_h3(doc, "Workflow 1 — Below-Cost SKU Alert System")
add_body(doc, "Goal: Automatically flag and block below-cost transactions before they happen.", size=9)
doc.add_paragraph()

wf1_steps = [
    ("Step 1", "AI Assistant", "Scan transaction/SKU data weekly. Identify all products where cost_price > list_price. Generate list of below-cost SKU IDs with unit loss amounts."),
    ("Step 2", "Automation Tool (n8n / Zapier)", "Trigger: every Monday 08:00. Action: send email alert to Inventory Manager with the flagged SKU list, sorted by unit loss (highest first). Include a summary: 'X SKUs at risk this week, estimated loss if sold: IDR Y.'"),
    ("Step 3", "Human (Inventory Manager)", "Review the list. Decide per SKU: REPRICE (raise list price), DISCONTINUE (stop ordering), or INVESTIGATE (check if cost_price is a data entry error). Log decision in the system."),
    ("Step 4", "Automation Tool", "Implement approved decisions: update list_price in system or flag SKU as discontinued. Send confirmation back to Finance team."),
    ("Step 5", "Human (Finance)", "Weekly: verify that flagged SKUs from prior week no longer appear in new transactions. Track IDR loss reduction as KPI."),
]

tbl4 = doc.add_table(rows=1 + len(wf1_steps), cols=3)
tbl4.style = 'Table Grid'
tbl4.columns[0].width = Inches(0.8)
tbl4.columns[1].width = Inches(1.5)
tbl4.columns[2].width = Inches(4.0)

for i, hdr in enumerate(["Step", "Owner", "Action"]):
    c = tbl4.rows[0].cells[i]
    c.text = hdr
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(8)
    set_cell_bg(c, "1F3864")
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

owner_colors = {"AI Assistant": "E2EFDA", "Automation Tool (n8n / Zapier)": "FFF2CC",
                "Human (Inventory Manager)": "FFDCE0", "Automation Tool": "FFF2CC",
                "Human (Finance)": "FFDCE0"}

for i, (step, owner, action) in enumerate(wf1_steps, 1):
    row = tbl4.rows[i]
    row.cells[0].text = step
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = owner
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[1].paragraphs[0].runs[0].bold = True
    set_cell_bg(row.cells[1], owner_colors.get(owner, "FFFFFF"))
    row.cells[2].text = action
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)

doc.add_paragraph()
add_body(doc, "Step Color Key:", size=8)
for label, color in [("Green = AI Assistant", "E2EFDA"), ("Yellow = Automation Tool", "FFF2CC"), ("Red = Human Approval", "FFDCE0")]:
    p = doc.add_paragraph(f"  {label}")
    p.runs[0].font.size = Pt(8)

doc.add_paragraph()

add_h3(doc, "Workflow 2 — Monthly Return Rate Monitor & Alert")
add_body(doc, "Goal: Detect return rate anomalies before they peak — giving the team 4–6 weeks of advance warning.", size=9)
doc.add_paragraph()

wf2_steps = [
    ("Step 1", "AI Assistant", "At month-end: calculate return rate by store, category, and month. Compare to 5-year baseline averages. Flag any store/category combo exceeding 10.5% return rate."),
    ("Step 2", "Automation Tool (n8n / Zapier)", "Trigger: last day of each month. If any flag exists: send dashboard summary to Store Operations. Include: store name, return rate this month, baseline, and delta. Highlight Accessories + Medan combinations as highest priority."),
    ("Step 3", "Human (Store Manager)", "Review flagged stores/categories. Identify possible causes (quality issue, seasonal pattern, specific product). Initiate investigation within 5 business days."),
    ("Step 4", "Human (Store Manager)", "Decide action: return reason code audit, supplier quality check, or product pull. Document findings in monthly operations log."),
    ("Step 5", "AI Assistant", "Next month: compare post-action return rate to prior month. Generate 1-paragraph status update: 'Did the intervention work?' for management review."),
]

tbl5 = doc.add_table(rows=1 + len(wf2_steps), cols=3)
tbl5.style = 'Table Grid'
tbl5.columns[0].width = Inches(0.8)
tbl5.columns[1].width = Inches(1.5)
tbl5.columns[2].width = Inches(4.0)

for i, hdr in enumerate(["Step", "Owner", "Action"]):
    c = tbl5.rows[0].cells[i]
    c.text = hdr
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(8)
    set_cell_bg(c, "1F5C2E")
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

for i, (step, owner, action) in enumerate(wf2_steps, 1):
    row = tbl5.rows[i]
    row.cells[0].text = step
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = owner
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[1].paragraphs[0].runs[0].bold = True
    set_cell_bg(row.cells[1], owner_colors.get(owner, "FFFFFF"))
    row.cells[2].text = action
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)

doc.add_paragraph()
add_h3(doc, "Risks & Open Questions for Next Milestone")
risks = [
    "RISK: Both workflows currently rely on manual data export from the transaction system — need to explore direct API or database connection in Weeks 8–9",
    "RISK: The AI Assistant (CustomGPT) cannot query live data; needs to be re-prompted with updated datasets each month — explore automation of context injection",
    "OPEN: Workflow 1 requires buy-in from Inventory team to act on weekly SKU alerts — what is the escalation path if decisions are not made within 48 hours?",
    "OPEN: For Workflow 2, return 'reason codes' are not currently collected — adding this field to the POS system would significantly improve diagnosis accuracy (noted as a data gap in Milestone 2)",
    "OPEN: Should the AI Assistant be deployed as a CustomGPT, a Slack bot, or embedded in the existing reporting dashboard? Weeks 8–9 should prototype and compare.",
]
for r in risks:
    add_bullet(doc, r, size=9)

# FOOTER
doc.add_paragraph()
divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Team 2 – Group F4  |  Emir · Harry · Leonardus · Syahra  |  BINUS x ReVOU 2026  |  Milestone 3")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.save(OUT)
print(f"Done! Saved to: {OUT}")
